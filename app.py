# app.py
import os
import tempfile
import streamlit as st

# لازم تكون أول Streamlit command
st.set_page_config(page_title="Graduation Project Checker", layout="wide")

from core.extract import extract_docx, extract_pdf
from core.checks import run_checks
from core.llm import simple_summary
from core.report import to_json
from core.storage import save_report


st.title("🎓 Graduation Project Checker (PDF/DOCX)")
st.write("ارفع ملف مشروع التخرج (PDF أو Word) وسأفحصه حسب قالب الجامعة + أعطيك ملخص وتنبيهات.")

uploaded = st.file_uploader("Upload your project file", type=["pdf", "docx"])

if uploaded:
    suffix = ".pdf" if uploaded.name.lower().endswith(".pdf") else ".docx"

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded.read())
        tmp_path = tmp.name

    try:
        # 1) Extract
        if suffix == ".docx":
            doc = extract_docx(tmp_path)
        else:
            doc = extract_pdf(tmp_path)

        # 2) Rule-based checks
        results = run_checks(doc.raw_text)
        missing_titles = [r.title for r in results if not r.passed]

        # 3) Summary (fallback)
        summary = simple_summary(doc.raw_text)

        # 5) Build report
        report = to_json(results, summary)

        # 6) DOCX Formatting Checks (Word only)
        format_issues = []
        if suffix == ".docx":
            from docx import Document
            from core.format_checks import check_abstract_format, check_captions

            docx_obj = Document(tmp_path)
            format_issues += check_abstract_format(docx_obj.paragraphs)
            format_issues += check_captions(docx_obj.paragraphs)

        report["format_issues"] = format_issues

        # 7) 💾 Save report
        saved_path = save_report(report, uploaded.name)
        st.success(f"تم حفظ التقرير: {saved_path}")

        # 8) UI
        col1, col2 = st.columns([1, 1])

        with col1:
            st.subheader("✅ النتيجة العامة")
            st.metric("Compliance Score", f"{report['score']}%")

            st.subheader("🧾 Summary (ملخص الفكرة)")
            st.write(
                report["summary"]
                if report["summary"]
                else "لم أستطع توليد ملخص واضح من الملف."
            )

            st.subheader("🤖 LLM Feedback (ملخص + ملاحظات ذكية)")
            st.info(
                "تم تعطيل الذكاء الاصطناعي في وضع العرض الآمن.\n"
                "يعتمد النظام حاليًا على فحص هيكلي وتقني قائم على القواعد الرسمية لقالب مشروع التخرج."
            )

            
            st.subheader("🧩 DOCX Formatting Checks")
            if suffix != ".docx":
                st.info("فحص التنسيق متاح لملفات Word فقط (DOCX).")
            elif len(report["format_issues"]) == 0:
                st.success("ما تم رصد مشاكل تنسيق أساسية في ملف الـDOCX ✅")
            else:
                for it in report["format_issues"]:
                    st.warning(f"**{it['what']}**\n\n**Fix:** {it['how']}")

        with col2:
            st.subheader("⚠️ Fix Suggestions (تنبيهات وإصلاحات)")
            if len(report["fixes"]) == 0:
                st.success("ملفك مستوفي الشروط الأساسية حسب القالب ✅")
            else:
                for f in report["fixes"]:
                    tag = "🔴" if f["priority"] == "high" else ("🟠" if f["priority"] == "medium" else "🟡")
                    st.warning(
                        f"{tag} {f['what']} — {f['details']}\n\n**What to do:** {f['how']}"
                    )

        st.divider()
        st.subheader("📋 Checklist (كل الفحوصات)")
        for c in report["checks"]:
            icon = "✅" if c["passed"] else "❌"
            st.write(f"{icon} **{c['title']}** — {c['details']}")

        st.divider()
        st.download_button(
            "Download report as JSON",
            data=str(report).encode("utf-8"),
            file_name="report.json",
            mime="application/json"
        )

    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass
